"""
fill_brief.py  —  Loti Labs content brief generator

Reads a JSON data file and fills the Loti Labs Brief.docx template,
producing a formatted deliverable in the deliverables/ folder.

Usage:
    py fill_brief.py data.json

JSON schema: see data-schema.json next to this file.
"""

import copy
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

TEMPLATE     = Path(__file__).parent / "Loti Labs Brief.docx"
DELIVERABLES = Path(__file__).parent / "deliverables"

# ── low-level helpers ─────────────────────────────────────────────────────────

def para_replace(para, old, new):
    """Replace `old` with `new` inside a paragraph.

    Tries run-by-run first. If the placeholder is split across runs,
    rewrites the combined text into run[0] and clears the rest.
    """
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    if old in para.text:
        combined = para.text.replace(old, new)
        for i, run in enumerate(para.runs):
            run.text = combined if i == 0 else ""
        return True
    return False


def set_cell(cell, text):
    """Overwrite the text of a table cell's first paragraph."""
    para = cell.paragraphs[0]
    for i, run in enumerate(para.runs):
        run.text = text if i == 0 else ""
    if not para.runs:
        para.add_run(text)


def clone_para(ref_para, new_text):
    """Deep-copy a paragraph element, setting its text to new_text."""
    el = copy.deepcopy(ref_para._element)
    t_nodes = el.findall(".//" + qn("w:t"))
    for i, t in enumerate(t_nodes):
        t.text = new_text if i == 0 else ""
    return el


def remove_para(para):
    p = para._element
    if p.getparent() is not None:
        p.getparent().remove(p)


def find_table(doc, *header_keywords):
    """Return the first table whose header row contains all header_keywords."""
    for table in doc.tables:
        if table.rows:
            row_text = " ".join(c.text for c in table.rows[0].cells)
            if all(kw in row_text for kw in header_keywords):
                return table
    return None


# ── section fillers ───────────────────────────────────────────────────────────

def fill_cover(doc, d):
    # Cover is always the first table in the document
    cell = doc.tables[0].rows[0].cells[0]
    for para in cell.paragraphs:
        para_replace(para, "[keyword]", d["keyword"])
        para_replace(para, "[month]",   d["month"])
        para_replace(para, "[year]",    d["year"])


def fill_doc_info(doc, d):
    table = find_table(doc, "Field", "Details")
    if not table:
        return
    for row in table.rows:
        label = row.cells[0].text.strip()
        if "Primary URL" in label:
            set_cell(row.cells[1], d["primary_url"])
        elif "Secondary URL" in label:
            set_cell(row.cells[1], d["secondary_url"])
        elif "Tertiary" in label:
            if d.get("tertiary_url"):
                set_cell(row.cells[1], d["tertiary_url"])
            else:
                table._tbl.remove(row._tr)
        elif label == "Content Type":
            set_cell(row.cells[1], d["content_type"])
        elif label == "Strategic Purpose":
            set_cell(row.cells[1], d["strategic_purpose"])
        elif "Word Count" in label:
            set_cell(row.cells[1], d["word_count_range"])


def fill_keywords(doc, d):
    additional = d.get("additional_keywords", [])
    for para in doc.paragraphs:
        if "[primary keyword]" in para.text:
            para_replace(para, "[primary keyword]", d["primary_keyword"])
        elif "[additional keywords; one per item]" in para.text:
            if not additional:
                remove_para(para)
                continue
            para_replace(para, "[additional keywords; one per item]", additional[0])
            anchor = para._element
            for kw in additional[1:]:
                new_el = clone_para(para, kw)
                idx = list(anchor.getparent()).index(anchor)
                anchor.getparent().insert(idx + 1, new_el)
                anchor = new_el


def fill_intent_audience(doc, d):
    audiences = d.get("target_audiences", [])
    for para in doc.paragraphs:
        if "[primary intent]" in para.text:
            para_replace(para, "[primary intent]", d["primary_intent"])
        elif "[secondary intent]" in para.text:
            para_replace(para, "[secondary intent]", d["secondary_intent"])
        elif "[target audiences; bullet list]" in para.text:
            if not audiences:
                remove_para(para)
                continue
            para_replace(para, "[target audiences; bullet list]", audiences[0])
            anchor = para._element
            for aud in audiences[1:]:
                new_el = clone_para(para, aud)
                idx = list(anchor.getparent()).index(anchor)
                anchor.getparent().insert(idx + 1, new_el)
                anchor = new_el


def fill_competitor_table(doc, d):
    """Replace all template data rows with actual competitor rows."""
    table = find_table(doc, "Competitor", "Word Count", "What They Cover Well")
    if not table:
        return

    competitors = d.get("competitors", [])
    # Save formatting reference from first data row, then remove all data rows
    ref_tr = copy.deepcopy(table.rows[1]._tr)
    for row in table.rows[1:]:
        table._tbl.remove(row._tr)

    for comp in competitors:
        new_tr = copy.deepcopy(ref_tr)
        cells  = new_tr.findall(".//" + qn("w:tc"))
        values = [
            comp.get("url", ""),
            str(comp.get("word_count", "")),
            comp.get("covers_well", ""),
            comp.get("gaps", ""),
        ]
        for cell, val in zip(cells, values):
            t_nodes = cell.findall(".//" + qn("w:t"))
            if t_nodes:
                t_nodes[0].text = val
                for t in t_nodes[1:]:
                    t.text = ""
        table._tbl.append(new_tr)


def fill_serp_analysis(doc, d):
    # Each SERP bullet has run[0]=bold label, run[1]=" [placeholder]"
    mapping = {
        "[content type dominating SERP]": d.get("serp_content_type", ""),
        "[average word count]":           d.get("serp_avg_word_count", ""),
        "[competitor content gaps]":      d.get("serp_competitor_gaps", ""),
        "[domain types ranking]":         d.get("serp_domain_types", ""),
        "[tone of ranking content]":      d.get("serp_tone", ""),
        "[2026 SERP context]":            d.get("serp_context_2026", ""),
    }
    for para in doc.paragraphs:
        for ph, val in mapping.items():
            if ph in para.text:
                if len(para.runs) >= 2:
                    para.runs[1].text = para.runs[1].text.replace(ph, " " + val)
                else:
                    para_replace(para, ph, val)
                break


def fill_paa(doc, d):
    paa_markers  = {"[People Also Ask 1]", "[People Also Ask 2]",
                    "[People Also Ask 3]", "[…]", "[...]", "[People Also Ask n]"}
    paa_paras    = [p for p in doc.paragraphs
                    if any(m in p.text for m in paa_markers)]
    questions    = d.get("paa_questions", [])

    if not paa_paras:
        return

    ref_para  = paa_paras[0]
    parent    = ref_para._element.getparent()
    insert_at = list(parent).index(ref_para._element)

    # Remove all template PAA paragraphs
    for para in paa_paras:
        el = para._element
        if el.getparent() is not None:
            el.getparent().remove(el)

    # Insert actual questions at the same position
    for q in questions:
        parent.insert(insert_at, clone_para(ref_para, q))
        insert_at += 1


def fill_internal_links(doc, d):
    links = d.get("internal_links", [])
    for para in doc.paragraphs:
        if "[bullet list; all the provided links" in para.text:
            if not links:
                remove_para(para)
                break
            para_replace(para, para.text.strip(), links[0])
            anchor = para._element
            for link in links[1:]:
                new_el = clone_para(para, link)
                idx = list(anchor.getparent()).index(anchor)
                anchor.getparent().insert(idx + 1, new_el)
                anchor = new_el
            break


def fill_outline(doc, d):
    """Replace template outline bullets with actual outline items.

    Template format per line: run[0]=bold "H1:" / run[1]=" [H1 title]"
    """
    outline = d.get("outline", [])
    # Outline items are bullet-list paragraphs (not Heading style).
    # Section template headings share "[heading]" but have a Heading style.
    outline_paras = [p for p in doc.paragraphs
                     if ("[H1 title]" in p.text or "[heading]" in p.text)
                     and not p.style.name.startswith("Heading")]

    if not outline_paras:
        return

    ref_para  = outline_paras[0]
    parent    = ref_para._element.getparent()
    insert_at = list(parent).index(ref_para._element)

    for para in outline_paras:
        el = para._element
        if el.getparent() is not None:
            el.getparent().remove(el)

    for item in outline:
        new_el = copy.deepcopy(ref_para._element)
        runs   = new_el.findall(".//" + qn("w:r"))
        label  = item["level"] + ":"
        text   = " " + item["heading"]
        if len(runs) >= 2:
            t0 = runs[0].find(qn("w:t"))
            t1 = runs[1].find(qn("w:t"))
            if t0 is not None: t0.text = label
            if t1 is not None: t1.text = text
        elif len(runs) == 1:
            t0 = runs[0].find(qn("w:t"))
            if t0 is not None: t0.text = label + text
        parent.insert(insert_at, new_el)
        insert_at += 1


def fill_sections(doc, d):
    """Replace template section headings + bullets with actual guidance."""
    sections = d.get("sections", [])

    def is_template(para):
        t = para.text.strip()
        return (
            "Introduction ([suggested word count])" in t or
            ("[H2 or H3]" in t and "[heading]" in t) or
            "[what to cover here; bullet list]" in t or
            t in ("[...]", "[…]")
        )

    template_paras = [p for p in doc.paragraphs if is_template(p)]
    if not template_paras:
        return

    ref_heading = next(
        (p for p in template_paras if p.style.name.startswith("Heading")),
        template_paras[0]
    )
    ref_bullet = next(
        (p for p in template_paras if "[what to cover here; bullet list]" in p.text),
        None
    )

    parent    = template_paras[0]._element.getparent()
    insert_at = list(parent).index(template_paras[0]._element)

    for para in template_paras:
        el = para._element
        if el.getparent() is not None:
            el.getparent().remove(el)

    for section in sections:
        heading    = section["heading"]
        word_count = section.get("word_count", "")
        guidance   = section.get("guidance", [])
        full_label = f"{heading} ({word_count})" if word_count else heading

        # Heading paragraph
        new_h = copy.deepcopy(ref_heading._element)
        runs  = new_h.findall(".//" + qn("w:r"))
        if runs:
            t = runs[0].find(qn("w:t"))
            if t is not None:
                t.text = full_label
            for r in runs[1:]:
                t = r.find(qn("w:t"))
                if t is not None:
                    t.text = ""
        parent.insert(insert_at, new_h)
        insert_at += 1

        # Guidance bullet paragraphs
        if ref_bullet:
            for point in guidance:
                new_b = copy.deepcopy(ref_bullet._element)
                t_nodes = new_b.findall(".//" + qn("w:t"))
                if t_nodes:
                    t_nodes[0].text = point
                    for t in t_nodes[1:]:
                        t.text = ""
                parent.insert(insert_at, new_b)
                insert_at += 1


# ── entry point ───────────────────────────────────────────────────────────────

def fill_brief(data_path: str):
    with open(data_path, encoding="utf-8") as f:
        d = json.load(f)

    doc = Document(TEMPLATE)

    fill_cover(doc, d)
    fill_doc_info(doc, d)
    fill_keywords(doc, d)
    fill_intent_audience(doc, d)
    fill_competitor_table(doc, d)
    fill_serp_analysis(doc, d)
    fill_paa(doc, d)
    fill_internal_links(doc, d)
    fill_outline(doc, d)
    fill_sections(doc, d)

    DELIVERABLES.mkdir(exist_ok=True)
    slug = re.sub(r"[^\w\s-]", "", d["keyword"]).strip().lower()
    slug = re.sub(r"[\s_]+", "-", slug)
    out  = DELIVERABLES / f"{slug}-brief.docx"
    doc.save(out)
    print(f"Saved: {out}")
    return str(out)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: py fill_brief.py data.json")
        sys.exit(1)
    fill_brief(sys.argv[1])
